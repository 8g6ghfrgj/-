import os
import tempfile
import hashlib
import logging
from typing import List, Set, Dict, Tuple
from datetime import datetime

from telethon import TelegramClient
from telethon.tl.types import Message

from link_utils import URL_REGEX, is_valid_link_for_extraction

# ======================
# Logging
# ======================

logger = logging.getLogger(__name__)

# ======================
# File Processing Tracker
# ======================

_processed_files: Dict[str, Dict] = {}  # تتبع الملفات المعالجة
MAX_CACHE_SIZE = 1000  # الحد الأقصى للملفات في الكاش


# ======================
# File Hash and Identification
# ======================

def get_file_hash(file_path: str) -> str:
    """إنشاء بصمة فريدة للملف"""
    try:
        hasher = hashlib.md5()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b''):
                hasher.update(chunk)
        return hasher.hexdigest()
    except Exception as e:
        logger.error(f"Error calculating file hash: {e}")
        return ""

def get_file_identifier(message: Message) -> Tuple[str, str]:
    """إنشاء معرف فريد للملف"""
    if not message.file:
        return "", ""
    
    # استخدام معرف الملف من Telegram
    file_id = getattr(message.file, 'id', 0)
    file_size = getattr(message.file, 'size', 0)
    
    # إنشاء معرف فريد
    identifier = f"{file_id}_{file_size}"
    
    # إضافة اسم الملف إذا كان موجوداً
    if hasattr(message.file, 'name') and message.file.name:
        filename = message.file.name.lower()
        identifier = f"{filename}_{identifier}"
    
    return identifier, filename if 'filename' in locals() else ""

def is_file_already_processed(file_identifier: str, file_hash: str = "") -> bool:
    """التحقق إذا كان الملف قد تم معالجته مسبقاً"""
    if file_identifier in _processed_files:
        cached_file = _processed_files[file_identifier]
        
        # إذا كان هناك hash، تحقق منه
        if file_hash and cached_file.get('hash') == file_hash:
            return True
        
        # إذا لم يكن هناك hash، اعتبر أن الملف قد تمت معالجته
        return True
    
    return False


# ======================
# Cache Management
# ======================

def add_to_processed_cache(file_identifier: str, file_hash: str, links_found: int):
    """إضافة ملف إلى الكاش"""
    # تنظيف الكاش إذا تجاوز الحد
    if len(_processed_files) >= MAX_CACHE_SIZE:
        # حذف أقدم الملفات
        oldest_keys = sorted(
            _processed_files.keys(),
            key=lambda k: _processed_files[k].get('timestamp', 0)
        )[:100]  # حذف 100 ملف قديم
        for key in oldest_keys:
            del _processed_files[key]
    
    # إضافة الملف الجديد
    _processed_files[file_identifier] = {
        'hash': file_hash,
        'timestamp': datetime.now().timestamp(),
        'links_found': links_found
    }

def clear_file_cache():
    """مسح كاش الملفات"""
    global _processed_files
    _processed_files.clear()
    logger.info("File cache cleared")


# ======================
# File Download with Progress
# ======================

async def download_file_with_progress(
    client: TelegramClient,
    message: Message,
    output_path: str
) -> bool:
    """تحميل الملف مع تتبع التقدم"""
    try:
        # التحقق من حجم الملف
        file_size = getattr(message.file, 'size', 0)
        MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB الحد الأقصى
        
        if file_size > MAX_FILE_SIZE:
            logger.warning(f"File too large: {file_size/1024/1024:.2f}MB")
            return False
        
        # تحميل الملف
        download_path = await client.download_media(
            message,
            output_path,
            progress_callback=_download_progress_callback if logger.isEnabledFor(logging.INFO) else None
        )
        
        return download_path is not None and os.path.exists(output_path)
        
    except Exception as e:
        logger.error(f"Error downloading file: {e}")
        return False

def _download_progress_callback(current: int, total: int):
    """عرض تقدم التحميل"""
    if total > 0:
        percent = (current / total) * 100
        logger.info(f"Downloading: {percent:.1f}% ({current}/{total} bytes)")


# ======================
# Supported File Types
# ======================

SUPPORTED_EXTENSIONS = {
    '.pdf': 'application/pdf',
    '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    '.doc': 'application/msword',
    '.txt': 'text/plain',
    '.rtf': 'application/rtf',
    '.odt': 'application/vnd.oasis.opendocument.text'
}

def is_file_supported(filename: str, mime_type: str = "") -> bool:
    """التحقق إذا كان نوع الملف مدعوماً"""
    if not filename:
        return False
    
    ext = os.path.splitext(filename.lower())[1]
    
    # التحقق من الامتداد
    if ext in SUPPORTED_EXTENSIONS:
        return True
    
    # التحقق من MIME type
    if mime_type:
        for supported_mime in SUPPORTED_EXTENSIONS.values():
            if mime_type in supported_mime:
                return True
    
    return False


# ======================
# Link Extraction Functions
# ======================

def _extract_from_pdf(path: str) -> Set[str]:
    """استخراج الروابط من PDF"""
    links: Set[str] = set()
    
    try:
        # محاولة PyPDF2 أولاً
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(path)
            
            for page_num, page in enumerate(reader.pages, 1):
                try:
                    text = page.extract_text() or ""
                    found_links = URL_REGEX.findall(text)
                    
                    # تصفية الروابط
                    for link in found_links:
                        if is_valid_link_for_extraction(link):
                            links.add(link.strip())
                    
                except Exception as e:
                    logger.warning(f"Error extracting from PDF page {page_num}: {e}")
                    
        except ImportError:
            # استخدام pdfplumber كبديل
            try:
                import pdfplumber
                with pdfplumber.open(path) as pdf:
                    for page_num, page in enumerate(pdf.pages, 1):
                        try:
                            text = page.extract_text() or ""
                            found_links = URL_REGEX.findall(text)
                            
                            for link in found_links:
                                if is_valid_link_for_extraction(link):
                                    links.add(link.strip())
                                    
                        except Exception as e:
                            logger.warning(f"Error extracting from PDF page {page_num} with pdfplumber: {e}")
            except ImportError:
                logger.error("No PDF extraction library found. Install PyPDF2 or pdfplumber.")
                
    except Exception as e:
        logger.error(f"Error extracting from PDF {os.path.basename(path)}: {e}")
    
    return links


def _extract_from_docx(path: str) -> Set[str]:
    """استخراج الروابط من DOCX"""
    links: Set[str] = set()
    
    try:
        from docx import Document
        
        doc = Document(path)
        
        # استخراج من الفقرات
        for para in doc.paragraphs:
            if para.text:
                found_links = URL_REGEX.findall(para.text)
                for link in found_links:
                    if is_valid_link_for_extraction(link):
                        links.add(link.strip())
        
        # استخراج من الجداول
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        found_links = URL_REGEX.findall(cell.text)
                        for link in found_links:
                            if is_valid_link_for_extraction(link):
                                links.add(link.strip())
        
        # استخراج من التعليقات (إن وجدت)
        if hasattr(doc, 'comments'):
            for comment in doc.comments:
                if comment.text:
                    found_links = URL_REGEX.findall(comment.text)
                    for link in found_links:
                        if is_valid_link_for_extraction(link):
                            links.add(link.strip())
                            
    except ImportError:
        logger.error("python-docx not installed. Cannot extract from DOCX files.")
    except Exception as e:
        logger.error(f"Error extracting from DOCX {os.path.basename(path)}: {e}")
    
    return links


def _extract_from_txt(path: str) -> Set[str]:
    """استخراج الروابط من ملفات نصية"""
    links: Set[str] = set()
    
    try:
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
            found_links = URL_REGEX.findall(content)
            
            for link in found_links:
                if is_valid_link_for_extraction(link):
                    links.add(link.strip())
                    
    except Exception as e:
        logger.error(f"Error extracting from text file {os.path.basename(path)}: {e}")
    
    return links


def _extract_from_doc(path: str) -> Set[str]:
    """استخراج الروابط من ملفات DOC القديمة"""
    links: Set[str] = set()
    
    try:
        # محاولة تحويل DOC إلى نص
        import subprocess
        
        # استخدام antiword إذا كان متوفراً
        try:
            result = subprocess.run(
                ['antiword', path],
                capture_output=True,
                text=True,
                encoding='utf-8',
                errors='ignore'
            )
            
            if result.returncode == 0:
                content = result.stdout
                found_links = URL_REGEX.findall(content)
                
                for link in found_links:
                    if is_valid_link_for_extraction(link):
                        links.add(link.strip())
        except (FileNotFoundError, subprocess.SubprocessError):
            logger.warning("antiword not found. Cannot extract from .doc files.")
            
    except Exception as e:
        logger.error(f"Error extracting from DOC file {os.path.basename(path)}: {e}")
    
    return links


# ======================
# Main Extraction Function
# ======================

async def extract_links_from_file(
    client: TelegramClient,
    message: Message
) -> List[str]:
    """
    استخراج الروابط من الملفات المدعومة
    مع تجنب معالجة نفس الملف مرتين
    """
    if not message.file:
        return []
    
    # الحصول على معلومات الملف
    filename = message.file.name or "file"
    mime_type = message.file.mime_type or ""
    
    # التحقق من دعم نوع الملف
    if not is_file_supported(filename, mime_type):
        logger.debug(f"File type not supported: {filename}")
        return []
    
    # إنشاء معرف فريد للملف
    file_identifier, original_filename = get_file_identifier(message)
    
    if not file_identifier:
        logger.warning("Could not create file identifier")
        return []
    
    # التحقق إذا كان الملف قد تم معالجته
    if is_file_already_processed(file_identifier):
        logger.info(f"File already processed: {original_filename or filename}")
        return []
    
    links: Set[str] = set()
    file_hash = ""
    
    with tempfile.TemporaryDirectory() as tmpdir:
        temp_path = os.path.join(tmpdir, filename)
        
        # تحميل الملف
        logger.info(f"Downloading file: {filename}")
        download_success = await download_file_with_progress(client, message, temp_path)
        
        if not download_success:
            logger.error(f"Failed to download file: {filename}")
            return []
        
        # حساب بصمة الملف
        file_hash = get_file_hash(temp_path)
        
        # التحقق مرة أخرى باستخدام البصمة
        if file_hash and is_file_already_processed(file_identifier, file_hash):
            logger.info(f"File already processed (by hash): {filename}")
            return []
        
        # استخراج الروابط حسب نوع الملف
        ext = os.path.splitext(filename.lower())[1]
        
        if ext == '.pdf' or mime_type == 'application/pdf':
            links = _extract_from_pdf(temp_path)
            
        elif ext == '.docx' or mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            links = _extract_from_docx(temp_path)
            
        elif ext == '.doc' or mime_type == 'application/msword':
            links = _extract_from_doc(temp_path)
            
        elif ext == '.txt' or mime_type == 'text/plain':
            links = _extract_from_txt(temp_path)
        
        # إضافة الملف إلى الكاش
        add_to_processed_cache(file_identifier, file_hash, len(links))
        
        logger.info(f"Extracted {len(links)} links from file: {filename}")
    
    return list(links)


# ======================
# Batch Processing
# ======================

async def process_message_files(
    client: TelegramClient,
    message: Message
) -> Dict[str, List[str]]:
    """
    معالجة جميع الملفات في الرسالة
    """
    result = {
        "total_files": 0,
        "processed_files": 0,
        "total_links": 0,
        "files": []
    }
    
    if not message.file:
        return result
    
    result["total_files"] = 1
    
    try:
        links = await extract_links_from_file(client, message)
        
        if links:
            result["processed_files"] = 1
            result["total_links"] = len(links)
            
            filename = message.file.name or "file"
            result["files"].append({
                "filename": filename,
                "links_found": len(links),
                "links": links
            })
            
            logger.info(f"Processed file {filename}: found {len(links)} links")
            
    except Exception as e:
        logger.error(f"Error processing file in message: {e}")
    
    return result


# ======================
# Statistics and Reporting
# ======================

def get_file_processing_stats() -> Dict:
    """الحصول على إحصائيات معالجة الملفات"""
    stats = {
        "cache_size": len(_processed_files),
        "processed_files": list(_processed_files.keys())[:10],  # آخر 10 ملفات
        "total_links_extracted": sum(f.get('links_found', 0) for f in _processed_files.values())
    }
    
    # إحصائيات حسب نوع الملف
    file_types = {}
    for file_id in _processed_files:
        # محاولة تحديد نوع الملف من المعرف
        if '.pdf' in file_id.lower():
            file_types['pdf'] = file_types.get('pdf', 0) + 1
        elif '.docx' in file_id.lower():
            file_types['docx'] = file_types.get('docx', 0) + 1
        elif '.doc' in file_id.lower():
            file_types['doc'] = file_types.get('doc', 0) + 1
        elif '.txt' in file_id.lower():
            file_types['txt'] = file_types.get('txt', 0) + 1
    
    stats["file_types"] = file_types
    
    return stats
